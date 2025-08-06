# export_utils.py
import streamlit as st
from typing import Dict, List
import datetime
import json
from pathlib import Path
import tempfile
import zipfile

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.pdfgen import canvas
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

class ExportUtils:
    def __init__(self):
        """Initialize export utilities"""
        self.docx_available = DOCX_AVAILABLE
        self.pdf_available = PDF_AVAILABLE
        
        if not self.docx_available:
            st.warning("DOCX export not available - install python-docx")
        if not self.pdf_available:
            st.warning("PDF export not available - install reportlab")
    
    def export_note(self, note_data: Dict, format: str = 'markdown') -> str:
        """
        Export a single note in specified format
        
        Args:
            note_data: Note dictionary
            format: Export format ('markdown', 'txt', 'json', 'docx', 'pdf')
            
        Returns:
            Exported content as string or file path
        """
        try:
            if format.lower() == 'markdown':
                return self._export_markdown(note_data)
            elif format.lower() == 'txt':
                return self._export_txt(note_data)
            elif format.lower() == 'json':
                return self._export_json(note_data)
            elif format.lower() == 'docx' and self.docx_available:
                return self._export_docx(note_data)
            elif format.lower() == 'pdf' and self.pdf_available:
                return self._export_pdf(note_data)
            else:
                st.error(f"Unsupported export format: {format}")
                return self._export_txt(note_data)  # Fallback
                
        except Exception as e:
            st.error(f"Export failed: {str(e)}")
            return f"Export Error: {str(e)}"
    
    def export_multiple_notes(self, notes: List[Dict], format: str = 'markdown') -> str:
        """
        Export multiple notes as a single file
        
        Args:
            notes: List of note dictionaries
            format: Export format
            
        Returns:
            Combined exported content
        """
        try:
            if format.lower() == 'markdown':
                return self._export_multiple_markdown(notes)
            elif format.lower() == 'txt':
                return self._export_multiple_txt(notes)
            elif format.lower() == 'json':
                return json.dumps(notes, ensure_ascii=False, indent=2)
            else:
                # Fallback to text format
                return self._export_multiple_txt(notes)
                
        except Exception as e:
            st.error(f"Multiple notes export failed: {str(e)}")
            return f"Export Error: {str(e)}"
    
    def create_export_archive(self, notes: List[Dict], formats: List[str] = ['markdown', 'json']) -> bytes:
        """
        Create a ZIP archive containing exported notes in multiple formats
        
        Args:
            notes: List of note dictionaries
            formats: List of export formats to include
            
        Returns:
            ZIP archive as bytes
        """
        try:
            # Create temporary directory
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_path = Path(temp_dir)
                
                # Export individual notes in each format
                for i, note in enumerate(notes):
                    note_id = note.get('id', f'note_{i}')[:8]
                    timestamp = note.get('timestamp', '')[:10]
                    
                    for fmt in formats:
                        try:
                            content = self.export_note(note, fmt)
                            filename = f"{timestamp}_{note_id}.{fmt}"
                            
                            with open(temp_path / filename, 'w', encoding='utf-8') as f:
                                f.write(content)
                        except Exception as e:
                            st.warning(f"Failed to export note {note_id} as {fmt}: {str(e)}")
                
                # Create combined exports
                for fmt in formats:
                    try:
                        combined_content = self.export_multiple_notes(notes, fmt)
                        combined_filename = f"all_notes.{fmt}"
                        
                        with open(temp_path / combined_filename, 'w', encoding='utf-8') as f:
                            f.write(combined_content)
                    except Exception as e:
                        st.warning(f"Failed to create combined {fmt} export: {str(e)}")
                
                # Create ZIP archive
                zip_buffer = tempfile.NamedTemporaryFile(delete=False)
                with zipfile.ZipFile(zip_buffer.name, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                    for file_path in temp_path.rglob('*'):
                        if file_path.is_file():
                            zip_file.write(file_path, file_path.name)
                
                # Read ZIP content
                with open(zip_buffer.name, 'rb') as f:
                    zip_content = f.read()
                
                # Cleanup
                Path(zip_buffer.name).unlink()
                
                return zip_content
                
        except Exception as e:
            st.error(f"Archive creation failed: {str(e)}")
            return b""
    
    def _export_markdown(self, note_data: Dict) -> str:
        """Export note as Markdown"""
        try:
            md_content = []
            
            # Title
            title = f"Voice Note - {note_data.get('language', 'Unknown')}"
            md_content.append(f"# {title}\n")
            
            # Metadata
            md_content.append("## Metadata\n")
            md_content.append(f"- **Date:** {note_data.get('timestamp', 'Unknown')}")
            md_content.append(f"- **Language:** {note_data.get('language', 'Unknown')}")
            md_content.append(f"- **ID:** {note_data.get('id', 'Unknown')}")
            if note_data.get('audio_file'):
                md_content.append(f"- **Audio File:** {note_data.get('audio_file')}")
            md_content.append("")
            
            # Transcription
            if note_data.get('transcription'):
                md_content.append("## Transcription\n")
                md_content.append(note_data.get('transcription'))
                md_content.append("")
            
            # Summary
            if note_data.get('summary'):
                md_content.append("## Summary\n")
                md_content.append(note_data.get('summary'))
                md_content.append("")
            
            # Keywords
            if note_data.get('keywords'):
                md_content.append("## Keywords\n")
                keywords = note_data.get('keywords', [])
                md_content.append(", ".join(keywords))
                md_content.append("")
            
            # Tags
            if note_data.get('tags'):
                md_content.append("## Tags\n")
                tags = note_data.get('tags', [])
                md_content.append(" ".join([f"#{tag}" for tag in tags]))
                md_content.append("")
            
            return "\n".join(md_content)
            
        except Exception as e:
            return f"Markdown export error: {str(e)}"
    
    def _export_txt(self, note_data: Dict) -> str:
        """Export note as plain text"""
        try:
            txt_content = []
            
            # Header
            title = f"VOICE NOTE - {note_data.get('language', 'Unknown').upper()}"
            txt_content.append(title)
            txt_content.append("=" * len(title))
            txt_content.append("")
            
            # Metadata
            txt_content.append("METADATA:")
            txt_content.append(f"Date: {note_data.get('timestamp', 'Unknown')}")
            txt_content.append(f"Language: {note_data.get('language', 'Unknown')}")
            txt_content.append(f"ID: {note_data.get('id', 'Unknown')}")
            if note_data.get('audio_file'):
                txt_content.append(f"Audio File: {note_data.get('audio_file')}")
            txt_content.append("")
            
            # Transcription
            if note_data.get('transcription'):
                txt_content.append("TRANSCRIPTION:")
                txt_content.append("-" * 13)
                txt_content.append(note_data.get('transcription'))
                txt_content.append("")
            
            # Summary
            if note_data.get('summary'):
                txt_content.append("SUMMARY:")
                txt_content.append("-" * 8)
                txt_content.append(note_data.get('summary'))
                txt_content.append("")
            
            # Keywords
            if note_data.get('keywords'):
                txt_content.append("KEYWORDS:")
                txt_content.append("-" * 9)
                keywords = note_data.get('keywords', [])
                txt_content.append(", ".join(keywords))
                txt_content.append("")
            
            # Tags
            if note_data.get('tags'):
                txt_content.append("TAGS:")
                txt_content.append("-" * 5)
                tags = note_data.get('tags', [])
                txt_content.append(", ".join(tags))
                txt_content.append("")
            
            return "\n".join(txt_content)
            
        except Exception as e:
            return f"Text export error: {str(e)}"
    
    def _export_json(self, note_data: Dict) -> str:
        """Export note as JSON"""
        try:
            # Create a clean copy for export
            export_data = {
                "id": note_data.get('id'),
                "timestamp": note_data.get('timestamp'),
                "language": note_data.get('language'),
                "language_code": note_data.get('language_code'),
                "transcription": note_data.get('transcription'),
                "summary": note_data.get('summary'),
                "keywords": note_data.get('keywords', []),
                "tags": note_data.get('tags', []),
                "audio_file": note_data.get('audio_file'),
                "source": note_data.get('source', 'voice_note'),
                "exported": datetime.datetime.now().isoformat()
            }
            
            return json.dumps(export_data, ensure_ascii=False, indent=2)
            
        except Exception as e:
            return f'{{"error": "JSON export failed: {str(e)}"}}'
    
    def _export_docx(self, note_data: Dict) -> str:
        """Export note as DOCX (returns file path)"""
        if not self.docx_available:
            return "DOCX export not available"
        
        try:
            # Create document
            doc = Document()
            
            # Title
            title = f"Voice Note - {note_data.get('language', 'Unknown')}"
            doc.add_heading(title, 0)
            
            # Metadata
            doc.add_heading('Metadata', level=1)
            metadata_para = doc.add_paragraph()
            metadata_para.add_run(f"Date: {note_data.get('timestamp', 'Unknown')}\n").bold = True
            metadata_para.add_run(f"Language: {note_data.get('language', 'Unknown')}\n")
            metadata_para.add_run(f"ID: {note_data.get('id', 'Unknown')}\n")
            if note_data.get('audio_file'):
                metadata_para.add_run(f"Audio File: {note_data.get('audio_file')}\n")
            
            # Transcription
            if note_data.get('transcription'):
                doc.add_heading('Transcription', level=1)
                doc.add_paragraph(note_data.get('transcription'))
            
            # Summary
            if note_data.get('summary'):
                doc.add_heading('Summary', level=1)
                doc.add_paragraph(note_data.get('summary'))
            
            # Keywords
            if note_data.get('keywords'):
                doc.add_heading('Keywords', level=1)
                keywords = note_data.get('keywords', [])
                doc.add_paragraph(", ".join(keywords))
            
            # Tags
            if note_data.get('tags'):
                doc.add_heading('Tags', level=1)
                tags = note_data.get('tags', [])
                doc.add_paragraph(" ".join([f"#{tag}" for tag in tags]))
            
            # Save to temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
            doc.save(temp_file.name)
            
            return temp_file.name
            
        except Exception as e:
            return f"DOCX export error: {str(e)}"
    
    def _export_pdf(self, note_data: Dict) -> str:
        """Export note as PDF (returns file path)"""
        if not self.pdf_available:
            return "PDF export not available"
        
        try:
            # Create temporary file
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
            
            # Create PDF document
            doc = SimpleDocTemplate(temp_file.name, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title = f"Voice Note - {note_data.get('language', 'Unknown')}"
            story.append(Paragraph(title, styles['Title']))
            story.append(Spacer(1, 12))
            
            # Metadata
            story.append(Paragraph("Metadata", styles['Heading1']))
            metadata_text = f"""
            <b>Date:</b> {note_data.get('timestamp', 'Unknown')}<br/>
            <b>Language:</b> {note_data.get('language', 'Unknown')}<br/>
            <b>ID:</b> {note_data.get('id', 'Unknown')}<br/>
            """
            if note_data.get('audio_file'):
                metadata_text += f"<b>Audio File:</b> {note_data.get('audio_file')}<br/>"
            
            story.append(Paragraph(metadata_text, styles['Normal']))
            story.append(Spacer(1, 12))
            
            # Transcription
            if note_data.get('transcription'):
                story.append(Paragraph("Transcription", styles['Heading1']))
                story.append(Paragraph(note_data.get('transcription'), styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Summary
            if note_data.get('summary'):
                story.append(Paragraph("Summary", styles['Heading1']))
                story.append(Paragraph(note_data.get('summary'), styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Keywords
            if note_data.get('keywords'):
                story.append(Paragraph("Keywords", styles['Heading1']))
                keywords = note_data.get('keywords', [])
                story.append(Paragraph(", ".join(keywords), styles['Normal']))
                story.append(Spacer(1, 12))
            
            # Tags
            if note_data.get('tags'):
                story.append(Paragraph("Tags", styles['Heading1']))
                tags = note_data.get('tags', [])
                story.append(Paragraph(" ".join([f"#{tag}" for tag in tags]), styles['Normal']))
            
            # Build PDF
            doc.build(story)
            
            return temp_file.name
            
        except Exception as e:
            return f"PDF export error: {str(e)}"
    
    def _export_multiple_markdown(self, notes: List[Dict]) -> str:
        """Export multiple notes as combined Markdown"""
        try:
            md_content = []
            
            # Title
            md_content.append("# WhispNote Export")
            md_content.append(f"*Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n")
            
            # Table of contents
            md_content.append("## Table of Contents\n")
            for i, note in enumerate(notes, 1):
                title = f"{note.get('language', 'Unknown')} - {note.get('timestamp', '')[:10]}"
                md_content.append(f"{i}. [{title}](#{i}-{title.lower().replace(' ', '-').replace('/', '-')})")
            md_content.append("")
            
            # Individual notes
            for i, note in enumerate(notes, 1):
                title = f"{note.get('language', 'Unknown')} - {note.get('timestamp', '')[:10]}"
                md_content.append(f"## {i}. {title}")
                md_content.append("")
                
                # Add note content (excluding the main title)
                note_content = self._export_markdown(note)
                # Remove the first line (main title) and add the content
                note_lines = note_content.split('\n')[1:]  # Skip first title line
                md_content.extend(note_lines)
                md_content.append("\n---\n")  # Separator
            
            return "\n".join(md_content)
            
        except Exception as e:
            return f"Multiple markdown export error: {str(e)}"
    
    def _export_multiple_txt(self, notes: List[Dict]) -> str:
        """Export multiple notes as combined text"""
        try:
            txt_content = []
            
            # Header
            header = "WHISPNOTE EXPORT"
            txt_content.append(header)
            txt_content.append("=" * len(header))
            txt_content.append(f"Generated on: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            txt_content.append(f"Total Notes: {len(notes)}")
            txt_content.append("")
            
            # Individual notes
            for i, note in enumerate(notes, 1):
                title = f"NOTE {i}: {note.get('language', 'Unknown')} - {note.get('timestamp', '')[:10]}"
                txt_content.append(title)
                txt_content.append("=" * len(title))
                
                # Add note content (excluding the main title)
                note_content = self._export_txt(note)
                # Remove the first few lines (main title and separator) and add the content
                note_lines = note_content.split('\n')[3:]  # Skip title and separator
                txt_content.extend(note_lines)
                txt_content.append("\n" + "-" * 80 + "\n")  # Separator
            
            return "\n".join(txt_content)
            
        except Exception as e:
            return f"Multiple text export error: {str(e)}"
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported export formats"""
        formats = ['markdown', 'txt', 'json']
        
        if self.docx_available:
            formats.append('docx')
        if self.pdf_available:
            formats.append('pdf')
            
        return formats
    
    def format_filename(self, note_data: Dict, format: str) -> str:
        """Generate appropriate filename for exported note"""
        try:
            timestamp = note_data.get('timestamp', '')[:10].replace('-', '')
            language = note_data.get('language', 'unknown').lower().replace(' ', '_')
            note_id = note_data.get('id', 'unknown')[:8]
            
            filename = f"{timestamp}_{language}_{note_id}.{format}"
            return filename
            
        except Exception:
            return f"note_export.{format}"